"""
Multi-Format Document Extractor

Handles:
- PDF slides (pypdf)
- DOCX documents (python-docx)
- Automatic format detection
- Structure preservation

For DSA-RAG-FEEIT thesis project
"""

from pypdf import PdfReader
from pathlib import Path
from typing import List, Dict, Optional, Union
import re


class MultiFormatExtractor:
    """Extract text from multiple document formats"""
    
    def __init__(self):
        self.issues = []
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check if required libraries are installed"""
        try:
            import docx
            self.has_docx = True
        except ImportError:
            self.has_docx = False
            print("⚠️  python-docx not installed. Install with: pip install python-docx")
    
    def extract_document(self, file_path: str) -> List[Dict]:
        """
        Auto-detect format and extract content.
        
        Args:
            file_path: Path to document
            
        Returns:
            List of extracted document chunks with metadata
        """
        path = Path(file_path)
        
        if not path.exists():
            return [{"error": f"File not found: {file_path}"}]
        
        # Detect format by extension
        extension = path.suffix.lower()
        
        if extension == '.pdf':
            return self._extract_pdf(file_path)
        elif extension in ['.docx', '.doc']:
            return self._extract_docx(file_path)
        else:
            return [{"error": f"Unsupported format: {extension}"}]
    
    def _extract_pdf(self, pdf_path: str) -> List[Dict]:
        """Extract text from PDF (slides or documents)"""
        try:
            reader = PdfReader(pdf_path)
        except Exception as e:
            return [{"error": f"Cannot read PDF: {str(e)}", "source": Path(pdf_path).name}]
        
        pages = []
        source_name = Path(pdf_path).name
        
        for i, page in enumerate(reader.pages):
            page_data = self._extract_pdf_page(page, source_name, i + 1)
            pages.append(page_data)
        
        return pages
    
    def _extract_docx(self, docx_path: str) -> List[Dict]:
        """Extract text from Word document with structure preservation"""
        if not self.has_docx:
            return [{"error": "python-docx not installed", "source": Path(docx_path).name}]
        
        try:
            import docx
            doc = docx.Document(docx_path)
        except Exception as e:
            return [{"error": f"Cannot read DOCX: {str(e)}", "source": Path(docx_path).name}]
        
        source_name = Path(docx_path).name
        
        # Extract paragraphs with style information
        content_blocks = []
        current_section = None
        section_content = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            if not text:
                continue
            
            # Detect headers (usually have heading styles or are short + bold)
            is_header = (
                para.style.name.startswith('Heading') or
                (len(text) < 60 and self._is_likely_header(text))
            )
            
            if is_header:
                # Save previous section
                if current_section and section_content:
                    content_blocks.append({
                        "type": "section",
                        "header": current_section,
                        "content": '\n'.join(section_content),
                        "source": source_name
                    })
                
                # Start new section
                current_section = text
                section_content = []
            else:
                section_content.append(text)
        
        # Add final section
        if current_section and section_content:
            content_blocks.append({
                "type": "section",
                "header": current_section,
                "content": '\n'.join(section_content),
                "source": source_name
            })
        
        # If no sections detected, treat as single document
        if not content_blocks:
            all_text = '\n'.join(para.text.strip() for para in doc.paragraphs if para.text.strip())
            content_blocks = [{
                "type": "document",
                "header": None,
                "content": all_text,
                "source": source_name
            }]
        
        # Add metadata
        for block in content_blocks:
            block["format"] = "docx"
            block["char_count"] = len(block["content"])
            block["has_code"] = self._has_code(block["content"])
            block["issues"] = self._detect_issues(block["content"])
        
        return content_blocks
    
    def _extract_pdf_page(self, page, source: str, page_num: int) -> Dict:
        """Extract text from single PDF page"""
        try:
            text = page.extract_text()
        except Exception as e:
            return {
                "source": source,
                "page_number": page_num,
                "text": "",
                "format": "pdf",
                "issues": [f"Extraction failed: {str(e)}"]
            }
        
        if not text:
            return {
                "source": source,
                "page_number": page_num,
                "text": "",
                "format": "pdf",
                "issues": ["Empty page - no text extracted"]
            }
        
        # Clean text
        text = self._clean_text(text)
        
        return {
            "source": source,
            "page_number": page_num,
            "text": text,
            "format": "pdf",
            "char_count": len(text),
            "has_code": self._has_code(text),
            "issues": self._detect_issues(text)
        }
    
    def _is_likely_header(self, text: str) -> bool:
        """Heuristic to detect headers in DOCX"""
        # Common header patterns
        header_indicators = [
            text.isupper(),  # ALL CAPS
            text.endswith(':'),  # Section header with colon
            len(text.split()) <= 6,  # Very short
            bool(re.match(r'^\d+\.?\s+[А-Яа-яA-Za-z]', text)),  # Starts with number
            text.startswith('•'),  # Bullet point header
        ]
        
        return sum(header_indicators) >= 2
    
    def _clean_text(self, text: str) -> str:
        """Clean extracted text while preserving structure"""
        # Remove null bytes
        text = text.replace('\x00', '')
        
        # Normalize whitespace but preserve newlines
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = ' '.join(line.split())  # Collapse spaces
            if line:
                cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)
    
    def _has_code(self, text: str) -> bool:
        """Detect if text contains code"""
        code_indicators = [
            r'\bfor\s*\(',
            r'\bwhile\s*\(',
            r'\bif\s*\(',
            r'\{[^}]{5,}\}',
            r'int\s+\w+\s*[=;]',
            r'void\s+\w+\s*\(',
            r'return\s+\w+;',
            r'\w+\[.*\]\s*=',
            r'public\s+class',
        ]
        
        return any(re.search(pattern, text) for pattern in code_indicators)
    
    def _detect_issues(self, text: str) -> List[str]:
        """Detect potential extraction issues"""
        issues = []
        
        if len(text) < 50:
            issues.append(f"Very short ({len(text)} chars)")
        
        if '\ufffd' in text:
            issues.append("Encoding errors detected")
        
        if self._has_code(text):
            open_braces = text.count('{')
            close_braces = text.count('}')
            if open_braces != close_braces:
                issues.append(f"Unbalanced braces ({open_braces}/{close_braces})")
        
        return issues
    
    def extract_batch(self, file_paths: List[str]) -> List[Dict]:
        """Extract from multiple files"""
        all_docs = []
        
        for path in file_paths:
            docs = self.extract_document(path)
            all_docs.extend(docs)
        
        return all_docs
    
    def get_extraction_report(self, documents: List[Dict]) -> Dict:
        """Generate extraction statistics"""
        total = len(documents)
        
        if total == 0:
            return {"error": "No documents extracted"}
        
        # Count by format
        from collections import Counter
        formats = Counter(d.get('format') for d in documents)
        
        # Issues
        docs_with_issues = sum(1 for d in documents if d.get('issues'))
        
        # Sizes
        char_counts = [d.get('char_count', 0) for d in documents if d.get('char_count')]
        avg_chars = sum(char_counts) / len(char_counts) if char_counts else 0
        
        return {
            "total_documents": total,
            "by_format": dict(formats),
            "documents_with_issues": docs_with_issues,
            "avg_chars_per_doc": int(avg_chars),
            "code_containing": sum(1 for d in documents if d.get('has_code')),
        }


# Convenience functions
def extract_pdf(pdf_path: str) -> List[Dict]:
    """Extract from PDF"""
    extractor = MultiFormatExtractor()
    return extractor.extract_document(pdf_path)


def extract_docx(docx_path: str) -> List[Dict]:
    """Extract from DOCX"""
    extractor = MultiFormatExtractor()
    return extractor.extract_document(docx_path)


if __name__ == "__main__":
    print("Multi-Format Extractor")
    print("Import and use: extract_pdf() or extract_docx()")
